from flask import Flask, render_template, request, jsonify, send_file, url_for, send_from_directory
import os
import json
import requests
import docx
from docx.shared import RGBColor
import io
import tempfile
import logging
import uuid
import shutil
import re
import time
import asyncio
import aiohttp
import aiofiles
import mimetypes
import zipfile
from firecrawl import AsyncFirecrawlApp
from urllib.parse import urljoin, urlparse
from werkzeug.utils import secure_filename
from PIL import Image
import datetime
import hashlib
import atexit
import signal
import pytesseract
import langdetect

# Clé API Firecrawl par défaut
FIRECRAWL_API_KEY = 'fc-4e2ef9d083654a49ab17a5dc27888c03'

# --- Serve .webp with correct MIME ---
mimetypes.add_type('image/webp', '.webp')

app = Flask(__name__)
app.config['SECRET_KEY'] = 'webtoon-editor-secret-key'
app.config['UPLOAD_FOLDER'] = os.path.join('static', 'uploads')
app.config['MAX_CONTENT_LENGTH'] = 100 * 1024 * 1024  # Augmenté à 100 MB max upload
logging.basicConfig(level=logging.INFO)

# Ensure the data and uploads directories exist
if not os.path.exists('data'):
    os.makedirs('data')
    
if not os.path.exists(app.config['UPLOAD_FOLDER']):
    os.makedirs(app.config['UPLOAD_FOLDER'])

# File to store project data
DATA_FILE = 'data/project_data.json'

# Initialize default data
def get_default_data():
    return {
        'blocks': [],
        'blockTypes': ['HB', 'B', 'DB', 'C', 'HC'],
        'api_keys': {
            'gemini': '',
            'firecrawl': ''
        },
        'glossary': []  # <-- AJOUT: Glossaire par défaut
    }

# Load project data
def load_data():
    if os.path.exists(DATA_FILE):
        try:
            with open(DATA_FILE, 'r', encoding='utf-8') as f:
                data = json.load(f)
                # Assurer la rétrocompatibilité pour les anciens projets
                data.setdefault('glossary', []) 
                return data
        except Exception as e:
            logging.error(f"Error loading data: {str(e)}")
            return get_default_data()
    return get_default_data()

# Save project data
def save_data(data):
    with open(DATA_FILE, 'w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=2)

# Verify if response contains a real image
def is_real_image(resp):
    """Check if the response contains a real image"""
    ct = resp.headers.get("Content-Type", "")
    if not ct.startswith("image/"):
        return False
    # 12 bytes "RIFF????WEBP" for WebP files
    if "webp" in ct:
        return resp.content.startswith(b"RIFF") and b"WEBP" in resp.content[:12]
    return True

@app.route('/')
def index():
    return render_template('index.html')

# Serve uploads explicitly
@app.route('/uploads/<path:fname>')
def serve_upload(fname):
    return send_from_directory(app.config['UPLOAD_FOLDER'], fname)

@app.route('/api/load', methods=['GET'])
def api_load():
    return jsonify(load_data())

@app.route('/api/save', methods=['POST'])
def api_save():
    data = request.json
    if data is None:
        return jsonify({"error": "Invalid JSON"}), 400
    
    current_data = load_data()
    current_data['blocks'] = data.get('blocks', current_data['blocks'])
    current_data['blockTypes'] = data.get('blockTypes', current_data['blockTypes'])
    # <-- AJOUT: Sauvegarde du glossaire
    current_data['glossary'] = data.get('glossary', current_data.get('glossary', []))
    
    save_data(current_data)
    return jsonify({"status": "success"})


# ----------  API KEYS ----------
@app.route('/api/get-keys')
def get_keys():
    return jsonify(load_data().get('api_keys', {}))

@app.route('/api/set-keys', methods=['POST'])
def set_keys():
    incoming = request.json or {}
    data = load_data()
    
    if 'api_keys' not in data:
        data['api_keys'] = {'gemini': '', 'firecrawl': ''}
        
    data['api_keys'].update({
        'gemini': incoming.get('gemini', data['api_keys'].get('gemini', '')),
        'firecrawl': incoming.get('firecrawl', data['api_keys'].get('firecrawl', ''))
    })
    save_data(data)
    return jsonify({"status": "ok"})

@app.route('/api/extract-webtoon', methods=['POST'])
def api_extract_webtoon():
    """Extract images from a webtoon URL"""
    data = request.json
    if data is None:
        return jsonify({"error": "Invalid JSON"}), 400
    
    url = data.get('url', '')
    if not url:
        return jsonify({"error": "URL is required"}), 400
    
    try:
        # Create a folder for this extraction
        extraction_id = str(uuid.uuid4())[:8]
        extraction_folder = os.path.join(app.config['UPLOAD_FOLDER'], f'webtoon_{extraction_id}')
        os.makedirs(extraction_folder, exist_ok=True)
        
        # Extract images
        images = extract_webtoon_images(url, extraction_folder)
        
        if not images:
            return jsonify({"error": "No images found or could not access the website"}), 404
        
        # Return image URLs and info
        return jsonify({
            "status": "success",
            "extraction_id": extraction_id,
            "images": images
        })
    except Exception as e:
        logging.error(f"Extraction error: {str(e)}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/extract-firecrawl', methods=['POST'])
def api_extract_firecrawl():
    """Extract images from a webtoon URL using Firecrawl"""
    data = request.json
    if data is None:
        return jsonify({"error": "Invalid JSON"}), 400
    
    url = data.get('url', '')
    api_key = data.get('api_key') or load_data().get('api_keys', {}).get('firecrawl', '')
    
    if not url:
        return jsonify({"error": "URL is required"}), 400
    
    if not api_key:
        return jsonify({"error": "Firecrawl API key missing"}), 400
    
    try:
        # Create a folder for this extraction
        extraction_id = str(uuid.uuid4())[:8]
        extraction_folder = os.path.join(app.config['UPLOAD_FOLDER'], f'webtoon_{extraction_id}')
        os.makedirs(extraction_folder, exist_ok=True)
        
        # Run the async extraction in a synchronous context
        images = asyncio.run(extract_firecrawl_images(url, api_key, extraction_folder))
        
        if not images:
            return jsonify({"error": "No images found or could not access the website"}), 404
        
        # Return image URLs and info
        return jsonify({
            "status": "success",
            "extraction_id": extraction_id,
            "images": images
        })
    except Exception as e:
        logging.error(f"Firecrawl extraction error: {str(e)}")
        return jsonify({"error": str(e)}), 500

def extract_webtoon_images(url, save_folder):
    """Extract images from a webtoon URL and save them to the specified folder using regex"""
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36',
        'Referer': url,                               # Add referer
        'Accept': 'image/avif,image/webp,*/*'         # Add modern accept header
    }
    
    try:
        # Get the webpage content
        response = requests.get(url, headers=headers, timeout=20)
        response.raise_for_status()
        
        html_content = response.text
        
        # Use regex to find image URLs
        # This pattern matches common image URLs in HTML
        img_patterns = [
            r'<img[^>]+src=["\'](https?://[^"\']+\.(jpg|jpeg|png|gif|webp))["\']',  # Standard img src
            r'<img[^>]+data-src=["\'](https?://[^"\']+\.(jpg|jpeg|png|gif|webp))["\']',  # Lazy loaded images
            r'<img[^>]+data-original=["\'](https?://[^"\']+\.(jpg|jpeg|png|gif|webp))["\']',  # Another lazy load pattern
            r'content=["\'](https?://[^"\']+\.(jpg|jpeg|png|gif|webp))["\']',  # Meta images
            r'background-image:\s*url\(["\']?(https?://[^"\']+\.(jpg|jpeg|png|gif|webp))["\']?\)',  # CSS background images
        ]
        
        # For asuracomic.net specifically
        if 'asuracomic.net' in url:
            logging.info("Detected asuracomic.net, using specific patterns")
            img_patterns.append(r'<div class="page-break[^>]*>.*?<img[^>]+src=["\'](https?://[^"\']+\.(jpg|jpeg|png|gif|webp))["\']')
            img_patterns.append(r'<div class="reading-content[^>]*>.*?<img[^>]+src=["\'](https?://[^"\']+\.(jpg|jpeg|png|gif|webp))["\']')
        
        all_image_urls = []
        for pattern in img_patterns:
            matches = re.findall(pattern, html_content)
            for match in matches:
                # The first group contains the URL
                if isinstance(match, tuple):
                    img_url = match[0]
                else:
                    img_url = match
                
                if img_url not in all_image_urls:
                    all_image_urls.append(img_url)
        
        logging.info(f"Found {len(all_image_urls)} potential image URLs")
        
        # Filter out small images and icons (usually width/height in attributes or URL contains 'icon', 'logo', etc.)
        filtered_urls = []
        for img_url in all_image_urls:
            if not any(x in img_url.lower() for x in ['icon', 'logo', 'banner', 'button', 'thumbnail', 'avatar']):
                filtered_urls.append(img_url)
        
        logging.info(f"After filtering: {len(filtered_urls)} image URLs")
        
        images = []
        for i, img_url in enumerate(filtered_urls):
            try:
                # Download the image
                img_response = requests.get(img_url, headers=headers, timeout=20)
                img_response.raise_for_status()
                
                # Check if it's a real image
                if not is_real_image(img_response):
                    logging.warning(f"Skipped non-image {img_url}")
                    continue
                
                # Determine file extension
                content_type = img_response.headers.get('Content-Type', '')
                ext = '.jpg'  # Default extension
                if 'image/png' in content_type:
                    ext = '.png'
                elif 'image/gif' in content_type:
                    ext = '.gif'
                elif 'image/webp' in content_type:
                    ext = '.webp'
                
                # Save the image
                img_filename = f'image_{i+1:03d}{ext}'
                img_path = os.path.join(save_folder, img_filename)
                
                # For WebP files, consider converting to PNG if needed
                if ext == '.webp':
                    try:
                        # Convert WebP -> PNG et ne conserver que le PNG pour éviter les doublons
                        im = Image.open(io.BytesIO(img_response.content)).convert("RGB")
                        png_filename = f'image_{i+1:03d}.png'
                        png_path = os.path.join(save_folder, png_filename)
                        im.save(png_path, "PNG", optimize=True)
                        img_filename_to_use = png_filename  # On utilisera la version PNG
                    except Exception as e:
                        logging.error(f"WEBP convert error {img_url}: {e}")
                        continue
                else:
                    with open(img_path, 'wb') as f:
                        f.write(img_response.content)
                    img_filename_to_use = img_filename
                
                # Create URL for the saved image using the new serve_upload route
                folder_id = os.path.basename(save_folder)
                img_url_path = url_for('serve_upload', fname=f'{folder_id}/{img_filename_to_use}')
                
                images.append({
                    'filename': img_filename_to_use,
                    'url': img_url_path,
                    'path': img_path
                })
                
                logging.info(f"Successfully saved image: {img_filename_to_use}")
                
                # Be nice to the server
                time.sleep(0.5)
                
            except Exception as e:
                logging.error(f"Error downloading image {img_url}: {str(e)}")
        
        return images
        
    except Exception as e:
        logging.error(f"Error extracting images from {url}: {str(e)}")
        raise e

async def extract_firecrawl_images(url, api_key, save_folder):
    """Extract images from a webtoon URL using Firecrawl and save them to the specified folder"""
    logging.info(f"Starting Firecrawl extraction from {url}")
    
    try:
        # Initialize Firecrawl
        app = AsyncFirecrawlApp(api_key=api_key)
        
        # Scrape the URL to get image URLs
        logging.info("Launching page scraping...")
        
        # Utiliser plusieurs formats pour maximiser les chances de trouver des images
        try:
            scrape_response = await app.scrape_url(
                url=url,
                formats=['html', 'screenshot']  # Utiliser html et screenshot
            )
            
            # Check if we got a valid response
            if not scrape_response:
                logging.error("Error: Empty response from Firecrawl")
                return []

            logging.info("Scraping completed. Analyzing response...")
            
            # Liste pour stocker les URLs d'images
            image_urls = []
            
            # 1. Extraire les images du HTML
            if hasattr(scrape_response, 'html') and scrape_response.html:
                from bs4 import BeautifulSoup
                
                # Créer un document HTML analysable
                soup = BeautifulSoup(scrape_response.html, 'lxml')
                
                # Conteneurs typiques des commentaires et profils
                excluded_containers = [
                    '.comment', '.comments', '.comment-section', 
                    '.user-profile', '.user-avatar', '.avatar',
                    '#comments', '#comment-section',
                    '.social-share', '.share-buttons'
                ]
                
                # Marquer les images dans des conteneurs à exclure
                excluded_images = set()
                for container_selector in excluded_containers:
                    try:
                        for container in soup.select(container_selector):
                            for img in container.select('img'):
                                for attr in ['src', 'data-src', 'data-original']:
                                    if attr in img.attrs:
                                        img_url = img[attr]
                                        if isinstance(img_url, str) and 'http' in img_url:
                                            excluded_images.add(img_url)
                    except Exception as e:
                        logging.debug(f"Error filtering container {container_selector}: {str(e)}")
                
                # Trouver toutes les balises image
                all_images = []
                for img_tag in soup.select('img'):
                    # Chercher dans différents attributs où les URLs d'image peuvent se trouver
                    for attr in ['src', 'data-src', 'data-original', 'data-lazy-src']:
                        if attr in img_tag.attrs:
                            img_url = img_tag[attr]
                            if isinstance(img_url, str) and 'http' in img_url:
                                # Calculer les dimensions de l'image
                                width = img_tag.get('width', '0')
                                height = img_tag.get('height', '0')
                                try:
                                    width_val = int(str(width).replace('px', '')) if width else 0
                                    height_val = int(str(height).replace('px', '')) if height else 0
                                except ValueError:
                                    width_val, height_val = 0, 0
                                
                                all_images.append({
                                    'url': img_url,
                                    'width': width_val,
                                    'height': height_val,
                                    'alt': str(img_tag.get('alt', '')),
                                    'class': str(img_tag.get('class', ''))
                                })
                                break
                
                # Filtrer les images pertinentes
                filtered_images = []
                
                # Mots-clés à exclure (uniquement les plus évidents)
                exclude_keywords = [
                    'avatar', 'profile', 'user', 'gravatar', 'icon-',
                    'facebook', 'twitter', 'pinterest', 'whatsapp', 'instagram'
                ]
                
                # Filtrage simple
                for img in all_images:
                    img_url = img['url']
                    
                    # Ignorer les images des conteneurs exclus (commentaires, etc.)
                    if img_url in excluded_images:
                        continue
                    
                    # Vérifier si c'est une petite icône sociale
                    url_lower = img_url.lower()
                    is_social_icon = any(social in url_lower for social in ['facebook', 'twitter', 'pinterest', 'instagram', 'whatsapp'])
                    
                    # Les très petites images sont probablement des icônes
                    is_tiny_image = (img['width'] > 0 and img['height'] > 0) and (img['width'] < 50 or img['height'] < 50)
                    
                    # Images de taille moyenne qui sont probablement des avatars
                    alt_text = img['alt'].lower()
                    is_profile_image = any(keyword in url_lower or keyword in alt_text for keyword in ['avatar', 'profile', 'user', 'gravatar'])
                    
                    # Exclure uniquement les petites images sociales et les avatars
                    if (is_social_icon and is_tiny_image) or (is_profile_image and img['width'] < 100 and img['height'] < 100):
                        continue
                    
                    # Exclure directement les .svg (souvent logos)
                    if img_url.lower().endswith('.svg'):
                        continue
                    
                    # Garder toutes les autres images
                    filtered_images.append(img_url)
                
                # Utiliser les images filtrées
                image_urls = filtered_images
                
                logging.info(f"Found {len(all_images)} total images, filtered to {len(image_urls)} relevant images")
            
            # 2. Si aucune image trouvée et qu'on a une capture d'écran, on la sauvegarde
            if len(image_urls) == 0 and hasattr(scrape_response, 'screenshot') and scrape_response.screenshot:
                screenshot_path = os.path.join(save_folder, 'page_screenshot.png')
                # S'assurer que screenshot est bien un objet bytes
                screenshot_data = scrape_response.screenshot
                if isinstance(screenshot_data, bytes):
                    with open(screenshot_path, 'wb') as f:
                        f.write(screenshot_data)
                    logging.info(f"Saved page screenshot to {screenshot_path}")
                else:
                    logging.error("Screenshot data is not in the expected format")
            
            # Log si aucune image trouvée
            if not image_urls:
                logging.error("No chapter images found in the scrape response.")
                return []
                
            logging.info(f"Found {len(image_urls)} chapter images in total.")
        
        except Exception as e:
            logging.error(f"Error during scraping: {str(e)}")
            return []

        # Download all images in parallel
        images = []
        async with aiohttp.ClientSession() as session:
            tasks = []
            for i, img_url in enumerate(image_urls):
                # Determine file extension from URL
                file_ext = os.path.splitext(img_url)[1]
                if not file_ext:  # If no extension, use .jpg as default
                    file_ext = '.jpg'
                
                # Create filename with proper numbering
                img_filename = f'image_{i+1:03d}{file_ext}'
                img_path = os.path.join(save_folder, img_filename)
                
                # Create task for downloading the image
                task = download_image(session, img_url, img_path, i, url)
                tasks.append(task)
            
            # Wait for all download tasks to complete
            download_results = await asyncio.gather(*tasks, return_exceptions=True)
            
            # Process download results
            for i, result in enumerate(download_results):
                if isinstance(result, Exception):
                    logging.error(f"Error downloading image {i+1}: {str(result)}")
                    continue
                
                # Skip None results (failed downloads)
                if result is None:
                    continue
                
                # Process successful downloads - we know result is a dict here
                if isinstance(result, dict):
                    img_filename = result.get('filename')
                    img_path = result.get('path')
                    
                    if img_filename and img_path:
                        # Create URL for the saved image using the new serve_upload route
                        folder_id = os.path.basename(save_folder)
                        img_url_path = url_for('serve_upload', fname=f'{folder_id}/{img_filename}')
                        
                        images.append({
                            'filename': img_filename,
                            'url': img_url_path,
                            'path': img_path
                        })
        
        return images
        
    except Exception as e:
        logging.error(f"Error extracting images with Firecrawl from {url}: {str(e)}")
        raise e

async def download_image(session, url, filepath, index, referer_url=None):
    """Download an image from a URL and save it to a file"""
    try:
        # Add proper headers for image requests
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36',
            'Accept': 'image/avif,image/webp,*/*'
        }
        
        # Add referer if provided
        if referer_url:
            headers['Referer'] = referer_url
            
        async with session.get(url, headers=headers) as response:
            if response.status == 200:
                content = await response.read()
                
                # Check if it's a real image by examining content type and data
                content_type = response.headers.get('Content-Type', '')
                is_image = content_type.startswith('image/')
                is_webp = 'webp' in content_type.lower()
                
                # For WebP, verify RIFF header
                if is_webp and (not content.startswith(b"RIFF") or b"WEBP" not in content[:12]):
                    logging.error(f"Invalid WebP image at {url}")
                    return None
                
                if not is_image:
                    logging.error(f"Not an image: {url} (Content-Type: {content_type})")
                    return None
                
                # Handle WebP images - convertir en PNG et ne pas conserver le .webp
                file_ext = os.path.splitext(filepath)[1].lower()
                if file_ext == '.webp':
                    try:
                        im = Image.open(io.BytesIO(content)).convert("RGB")
                        png_path = filepath.replace('.webp', '.png')
                        # Sauvegarde synchronously (Pillow requires file path or sync file object)
                        im.save(png_path, "PNG", optimize=True)
                        logging.info(f"Successfully downloaded and converted (PNG only): {png_path}")
                        saved_filename = os.path.basename(png_path)
                        saved_path = png_path
                    except Exception as e:
                        logging.error(f"WEBP convert error for {url}: {str(e)}")
                        return None
                else:
                    # Save non-WebP image
                    async with aiofiles.open(filepath, mode='wb') as f:
                        await f.write(content)
                    
                    logging.info(f"Successfully downloaded: {filepath}")
                    saved_filename = os.path.basename(filepath)
                    saved_path = filepath
                
                return {
                    'filename': saved_filename,
                    'path': saved_path,
                    'index': index
                }
            else:
                logging.error(f"Error {response.status} for URL: {url}")
                return None
    except Exception as e:
        logging.error(f"Failed to download {url}: {str(e)}")
        raise e

def cleanup_uploads_folder():
    uploads_dir = app.config['UPLOAD_FOLDER']
    if os.path.exists(uploads_dir):
        for item in os.listdir(uploads_dir):
            item_path = os.path.join(uploads_dir, item)
            try:
                if os.path.isdir(item_path):
                    shutil.rmtree(item_path)
                else:
                    os.remove(item_path)
            except Exception as e:
                logging.error(f"Erreur suppression {item_path}: {e}")

# Nettoyage au démarrage
cleanup_uploads_folder()

# Nettoyage à la fermeture
atexit.register(cleanup_uploads_folder)

def handle_exit(*args):
    cleanup_uploads_folder()
    os._exit(0)

signal.signal(signal.SIGINT, handle_exit)
signal.signal(signal.SIGTERM, handle_exit)

# Désactiver toute sauvegarde d'image hors import/export ZIP/DOCX
def is_upload_allowed():
    # Désactive toute sauvegarde d'image hors import/export
    return False

@app.route('/api/upload', methods=['POST'])
def api_upload():
    if not is_upload_allowed():
        return jsonify({"error": "Upload désactivé (gestion automatisée)"}), 403
    if 'file' not in request.files:
        return jsonify({"error": "No file part"}), 400
        
    file = request.files['file']
    
    if file.filename == '':
        return jsonify({"error": "No selected file"}), 400
        
    if file and file.filename:
        # Create a unique filename
        filename = secure_filename(file.filename)
        file_ext = os.path.splitext(filename)[1]
        unique_filename = f"{uuid.uuid4()}{file_ext}"
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], unique_filename)
        
        # Save the file
        file.save(file_path)
        
        # Return the path to access the file
        file_url = url_for('static', filename=f'uploads/{unique_filename}')
        
        return jsonify({
            "status": "success",
            "filename": filename,
            "path": file_path,
            "url": file_url
        })
    
    return jsonify({"error": "Failed to upload file"}), 500

@app.route('/api/cleanup', methods=['POST'])
def api_cleanup():
    """Cleanup uploads folder to keep only the most recent webtoon folder"""
    try:
        # Get uploads folder
        uploads_dir = app.config['UPLOAD_FOLDER']
        
        # Count files and directories before cleanup
        total_dirs_before = 0
        total_files_before = 0
        
        for root, dirs, files in os.walk(uploads_dir):
            total_dirs_before += len(dirs)
            total_files_before += len(files)
        
        # Get all webtoon directories
        webtoon_dirs = [d for d in os.listdir(uploads_dir) if d.startswith('webtoon_') and os.path.isdir(os.path.join(uploads_dir, d))]
        
        # Sort by modification time (newest first)
        webtoon_dirs.sort(key=lambda x: os.path.getmtime(os.path.join(uploads_dir, x)), reverse=True)
        
        # Keep only the most recent directory
        keep_folder = webtoon_dirs[0] if webtoon_dirs else None
        
        # Delete all other webtoon folders
        deleted_dirs = 0
        deleted_files = 0
        
        for item in os.listdir(uploads_dir):
            item_path = os.path.join(uploads_dir, item)
            
            # Only process webtoon directories
            if item.startswith('webtoon_') and os.path.isdir(item_path):
                if keep_folder and item == keep_folder:
                    # Skip the most recent folder
                    continue
                
                # Count files in this directory
                file_count = sum(len(files) for _, _, files in os.walk(item_path))
                deleted_files += file_count
                deleted_dirs += 1
                
                # Delete the directory
                shutil.rmtree(item_path)
                logging.info(f"Deleted directory: {item}")
        
        # Count files and directories after cleanup
        total_dirs_after = 0
        total_files_after = 0
        
        for root, dirs, files in os.walk(uploads_dir):
            total_dirs_after += len(dirs)
            total_files_after += len(files)
        
        # Return cleanup statistics
        return jsonify({
            "status": "success", 
            "message": f"Nettoyage terminé. {deleted_dirs} dossiers et {deleted_files} fichiers supprimés.",
            "stats": {
                "before": {"directories": total_dirs_before, "files": total_files_before},
                "after": {"directories": total_dirs_after, "files": total_files_after},
                "deleted": {"directories": deleted_dirs, "files": deleted_files}
            }
        })
    except Exception as e:
        logging.error(f"Cleanup error: {str(e)}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/spellcheck', methods=['POST'])
def api_spellcheck():
    data = request.json
    if data is None:
        return jsonify({"error": "Invalid JSON"}), 400
    
    blocks = data.get('blocks', [])
    
    if not blocks:
        return jsonify({"error": "No blocks to check"})
    
    all_errors = []
    
    try:
        # Appeler LanguageTool pour chaque bloc individuellement
        for idx, block in enumerate(blocks):
            if not block.get('content', '').strip():
                continue  # Ignorer les blocs vides
                
            # Appel à LanguageTool pour ce bloc uniquement
            response = requests.post(
                'https://api.languagetool.org/v2/check',
                data={
                    'text': block['content'],
                    'language': 'fr'
                },
                timeout=10  # Timeout pour éviter les blocages
            )
            
            if response.status_code == 200:
                data = response.json()
                
                # Traiter les matches pour ce bloc
                for match in data.get('matches', [])[:20]:  # Limiter à 20 erreurs par bloc
                    error_obj = {
                        'block_index': idx,
                        'error_start': match['offset'],
                        'error_end': match['offset'] + match.get('length', 0),
                        'message': match.get('message', ''),
                        'rule_id': match.get('rule', {}).get('id', ''),
                        'replacements': match.get('replacements', []),
                        'context': match.get('context', {}).get('text', ''),
                        'offset': match['offset'],
                        'length': match.get('length', 0)
                    }
                    all_errors.append(error_obj)
            else:
                logging.warning(f"LanguageTool error for block {idx}: {response.status_code}")
                
        # Also return the old format for backward compatibility
        total_errors = len(all_errors)
        if total_errors == 0:
            return jsonify({
                "report": "✅ Aucun problème",
                "errors": [],
                "total_errors": 0
            })
        
        # Créer le rapport texte pour la compatibilité
        report = f"Problèmes : {total_errors}\n\n"
        for i, error in enumerate(all_errors[:20]):  # Limiter à 20 dans le rapport
            report += f"→ {error['message']}\n"
        
        return jsonify({
            "report": report,
            "errors": all_errors,
            "total_errors": total_errors
        })
        
    except Exception as e:
        logging.error(f"Spellcheck error: {str(e)}")
        return jsonify({"error": "Network Error", "details": str(e)})

@app.route('/api/gemini', methods=['POST'])
def api_gemini():
    data = request.json
    if data is None:
        return jsonify({"error": "Invalid JSON"}), 400
    
    api_key = data.get('api_key', '')
    content = data.get('content', '')
    
    if not api_key or not content:
        return jsonify({"error": "Missing API key or content"})
    
    prompt = f"""
Réécris ou corrige en français.
Fournis 3 variantes concises, **sans parenthèses ni explication**, une par ligne :
« {content} »"""
    
    try:
        response = requests.post(
            f'https://generativelanguage.googleapis.com/v1beta/models/gemini-1.5-flash:generateContent?key={api_key}',
            json={
                'contents': [{'parts': [{'text': prompt}]}],
                'generationConfig': {'temperature': 0.8, 'candidateCount': 1}
            }
        )
        
        if response.status_code != 200:
            return jsonify({"error": response.json().get('error', {}).get('message', response.text)})
        
        data = response.json()
        raw_text = data.get('candidates', [{}])[0].get('content', {}).get('parts', [{}])[0].get('text', '')
        
        suggestions = []
        for line in raw_text.split('\n'):
            line = line.strip()
            if line:
                # Remove bullets, numbers, and explanations in parentheses
                line = line.lstrip('•-0123456789. \t')
                line = line.split('(', 1)[0].strip()
                if line:
                    suggestions.append(line)
        
        return jsonify({"suggestions": suggestions[:3]})
    except Exception as e:
        logging.error(f"Gemini API error: {str(e)}")
        return jsonify({"error": str(e)})

@app.route('/api/export-docx', methods=['POST'])
def export_docx():
    data = request.json
    if data is None:
        return jsonify({"error": "Invalid JSON"}), 400
    
    blocks = data.get('blocks', [])
    
    if not blocks:
        return jsonify({"error": "No blocks to export"})
    
    # Create document
    doc = docx.Document()
    
    # Add blocks in requested format
    for block in blocks:
        # Block name (type+number)
        doc.add_paragraph(f"{block['type']}{block['number']}")
        # Block content
        doc.add_paragraph(block['content'])
        # Double new line (empty paragraph)
        doc.add_paragraph("")
    
    # Save to temporary file
    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
        doc.save(tmp.name)
        tmp_path = tmp.name
    
    # Send file
    return send_file(
        tmp_path,
        as_attachment=True,
        download_name='script_webtoon.docx',
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

@app.route('/api/create-folder', methods=['POST'])
def api_create_folder():
    """Create a folder in the uploads directory"""
    try:
        data = request.json
        if not data or 'folderName' not in data:
            return jsonify({"error": "Folder name is required"}), 400
        
        folder_name = secure_filename(data['folderName'])
        folder_path = os.path.join(app.config['UPLOAD_FOLDER'], folder_name)
        
        # Create folder if it doesn't exist
        if not os.path.exists(folder_path):
            os.makedirs(folder_path)
            logging.info(f"Created folder: {folder_path}")
        
        return jsonify({
            "status": "success",
            "message": f"Folder {folder_name} created successfully",
            "path": folder_path
        })
    except Exception as e:
        logging.error(f"Error creating folder: {str(e)}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/upload-to-folder', methods=['POST'])
def api_upload_to_folder():
    if not is_upload_allowed():
        return jsonify({"error": "Upload désactivé (gestion automatisée)"}), 403
    if 'file' not in request.files:
        return jsonify({"error": "No file part"}), 400
        
    file = request.files['file']
    target_folder = request.form.get('targetFolder', '')
    
    if file.filename == '':
        return jsonify({"error": "No selected file"}), 400
        
    if file and file.filename:
        # Secure the filename
        filename = secure_filename(file.filename)
        
        # Ensure target folder exists and is within uploads directory
        target_folder = secure_filename(target_folder)
        folder_path = os.path.join(app.config['UPLOAD_FOLDER'], target_folder)
        
        if not os.path.exists(folder_path):
            os.makedirs(folder_path)
            logging.info(f"Created folder: {folder_path}")
        
        # Save the file to the target folder
        file_path = os.path.join(folder_path, filename)
        file.save(file_path)
        
        # Return the path to access the file
        file_url = url_for('serve_upload', fname=f'{target_folder}/{filename}')
        
        return jsonify({
            "status": "success",
            "filename": filename,
            "path": file_path,
            "url": file_url
        })
    
    return jsonify({"error": "Failed to upload file"}), 500

# ── helper ──────────────────────────────────────────────────────────────────────────
def latest_webtoon_folder() -> str | None:
    """Retourne le chemin du dossier webtoon_* le plus récent dans uploads/."""
    uploads = app.config['UPLOAD_FOLDER']
    candidates = [
        os.path.join(uploads, d) for d in os.listdir(uploads)
        if d.startswith("webtoon_") and os.path.isdir(os.path.join(uploads, d))
    ]
    if not candidates:
        return None
    return max(candidates, key=os.path.getmtime)

# ── nouvelle route d'export ZIP ────────────────────────────────────────────────
@app.route('/api/export-zip')
def api_export_zip():
    """
    Crée une archive ZIP contenant :
      • project_data.json   (blocs + clés + types)
      • /images/*           (toutes les images du dossier webtoon_* le plus récent)
      Les doublons d'images (même contenu) sont ignorés.
    """
    try:
        project_json = load_data()
        img_folder   = latest_webtoon_folder()          # peut être None

        # fichier temporaire .zip
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".zip")
        tmp.close()  # on le referme, zipfile va le rouvrir

        with zipfile.ZipFile(tmp.name, "w", zipfile.ZIP_DEFLATED) as z:
            # 1) JSON
            z.writestr("project_data.json", json.dumps(project_json, ensure_ascii=False, indent=2))

            # 2) Images (si dossier trouvé) avec déduplication par hash
            if img_folder:
                seen_hashes: set[str] = set()
                for root, _, files in os.walk(img_folder):
                    for f in files:
                        abs_path = os.path.join(root, f)
                        # calcul hash
                        try:
                            with open(abs_path, 'rb') as fp:
                                file_bytes = fp.read()
                                file_hash  = hashlib.sha256(file_bytes).hexdigest()
                        except Exception as e:
                            logging.error(f"Erreur lecture image {abs_path}: {e}")
                            continue

                        if file_hash in seen_hashes:
                            logging.info(f"Doublon ignoré: {f}")
                            continue  # skip duplicate content

                        seen_hashes.add(file_hash)
                        z.write(abs_path, arcname=os.path.join("images", f))

        return send_file(
            tmp.name,
            as_attachment=True,
            download_name="webtoon_export.zip",
            mimetype="application/zip"
        )
    except Exception as e:
        logging.error(f"ZIP export error: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/example-json')
def api_example_json():
    """Route pour télécharger un exemple de fichier project_data.json valide"""
    try:
        # Créer un exemple simple
        example_data = {
            "blocks": [
                {
                    "id": "1",
                    "type": "B",
                    "number": 1,
                    "content": "Exemple de bloc de texte"
                }
            ],
            "blockTypes": ["HB", "B", "DB", "C", "HC"],
            "exportVersion": "2.0",
            "exportDate": datetime.datetime.now().isoformat()
        }
        
        # Créer un fichier temporaire
        with tempfile.NamedTemporaryFile(delete=False, suffix=".json") as tmp:
            tmp.write(json.dumps(example_data, ensure_ascii=False, indent=2).encode('utf-8'))
            tmp_path = tmp.name
            
        return send_file(
            tmp_path,
            as_attachment=True,
            download_name="project_data.json",
            mimetype="application/json"
        )
    except Exception as e:
        logging.error(f"Erreur exemple JSON: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/example-zip')
def api_example_zip():
    """Route pour télécharger un exemple de fichier ZIP valide avec structure correcte"""
    try:
        # Créer un exemple simple de données de projet
        example_data = {
            "blocks": [
                {
                    "id": "1",
                    "type": "B",
                    "number": 1,
                    "content": "Exemple de bloc de texte"
                },
                {
                    "id": "2",
                    "type": "C",
                    "number": 1,
                    "content": "Ceci est un commentaire d'exemple"
                }
            ],
            "blockTypes": ["HB", "B", "DB", "C", "HC"],
            "exportVersion": "2.0",
            "exportDate": datetime.datetime.now().isoformat()
        }
        
        # Créer un fichier ZIP temporaire
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".zip")
        tmp.close()
        
        # Créer le zip avec la structure correcte
        with zipfile.ZipFile(tmp.name, "w", zipfile.ZIP_DEFLATED) as z:
            # 1. Ajouter le fichier project_data.json
            z.writestr("project_data.json", json.dumps(example_data, ensure_ascii=False, indent=2))
            
            # 2. Créer un dossier images/ vide (pas nécessaire, mais pour la démonstration)
            z.writestr("images/.keep", "")
        
        # Envoyer le fichier
        return send_file(
            tmp.name,
            as_attachment=True,
            download_name="exemple_import.zip",
            mimetype="application/zip"
        )
    except Exception as e:
        logging.error(f"Erreur création exemple ZIP: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/import-zip', methods=['POST'])
def api_import_zip():
    """
    Attend un fichier ZIP contenant :
      - project_data.json (n'importe où dans l'archive)
      - images/*  (facultatif, avec hiérarchie préservée)

    • Met à jour le fichier project_data.json sur le serveur
    • Copie les images dans un nouveau dossier uploads/webtoon_<id>
    • Renvoie les infos pour recharger l'UI
    """
    try:
        logging.info("=== Début de l'import ZIP ===")
        
        # Vérifier la présence du fichier
        if 'file' not in request.files:
            logging.error("Pas de partie 'file' dans la requête")
            return jsonify({"error": "No file part", "details": "Le formulaire doit contenir un champ 'file'"}), 400
        
        up = request.files['file']
        logging.info(f"Fichier reçu: {up.filename}, type MIME: {up.content_type}")
        
        if not up.filename:
            logging.error("Nom de fichier vide")
            return jsonify({"error": "Empty filename", "details": "Le nom du fichier est vide"}), 400
        
        # Vérifier le type de fichier
        allowed_extensions = {'.zip', '.wtoon'}
        if not any(up.filename.endswith(ext) for ext in allowed_extensions):
            logging.error(f"Extension non valide: {up.filename}")
            return jsonify({"error": "Need a .zip or .wtoon file"}), 400
    except Exception as e:
        # Assurer que même les erreurs inattendues renvoient du JSON valide
        logging.exception("Erreur inattendue lors de la validation du fichier")
        return jsonify({"error": "Erreur serveur", "details": str(e)}), 500

    try:
        with tempfile.TemporaryDirectory() as tmpdir:
            logging.info(f"Dossier temporaire créé: {tmpdir}")
            zpath = os.path.join(tmpdir, 'in.zip')
            up.save(zpath)
            logging.info(f"Fichier ZIP sauvegardé: {zpath}, taille: {os.path.getsize(zpath)} octets")

            try:
                with zipfile.ZipFile(zpath) as z:
                    # Log des fichiers contenus dans le ZIP
                    file_list = z.namelist()
                    logging.info(f"Contenu du ZIP: {len(file_list)} fichiers")
                    if len(file_list) > 0:
                        logging.info(f"Premiers fichiers: {', '.join(file_list[:5])}")
                    
                    # ----- trouver le JSON peu importe le dossier -------------
                    json_name = next((n for n in file_list 
                                      if os.path.basename(n).lower() == 'project_data.json'), None)
                    if not json_name:
                        logging.error("Aucun fichier project_data.json trouvé dans l'archive")
                        return jsonify({
                            "error": "project_data.json missing", 
                            "details": "Le fichier project_data.json est introuvable dans l'archive",
                            "files_found": file_list[:10] if len(file_list) <= 10 else file_list[:10] + ["..."]
                        }), 400
                    
                    logging.info(f"Fichier JSON trouvé: {json_name}")
                    
                    try:
                        # Tenter de lire le fichier JSON
                        json_content = z.read(json_name)
                        logging.info(f"Taille du JSON: {len(json_content)} octets")
                        
                        # Vérifier si le contenu ressemble à du HTML/XML (commence par <)
                        decoded_start = json_content[:100].decode('utf-8', errors='ignore').strip()
                        logging.info(f"Début du contenu: {decoded_start}")
                        if decoded_start.startswith('<'):
                            logging.error("Le fichier JSON contient du HTML/XML au lieu de JSON")
                            return jsonify({
                                "error": "HTML detected", 
                                "details": "Le fichier project_data.json contient du HTML au lieu de JSON. Vérifiez que vous avez bien exporté un fichier JSON valide.",
                                "content_preview": decoded_start
                            }), 400
                        
                        # Essayer différents encodages si UTF-8 échoue
                        proj_data = None
                        decode_errors = []
                        
                        for encoding in ['utf-8', 'utf-8-sig', 'latin1', 'cp1252']:
                            try:
                                decoded = json_content.decode(encoding)
                                proj_data = json.loads(decoded)
                                logging.info(f"Décodage réussi avec l'encodage: {encoding}")
                                break
                            except UnicodeDecodeError as ude:
                                decode_errors.append(f"{encoding}: {str(ude)}")
                                continue
                            except json.JSONDecodeError as jde:
                                decode_errors.append(f"{encoding}: {str(jde)}")
                                continue
                            
                        if proj_data is None:
                            logging.error(f"Échec du décodage avec tous les encodages: {decode_errors}")
                            return jsonify({
                                "error": "JSON decode error", 
                                "details": "Impossible de décoder le fichier JSON avec les encodages connus",
                                "decode_attempts": decode_errors,
                                "content_preview": decoded_start
                            }), 400
                            
                    except Exception as e:
                        logging.exception("Erreur lors de la lecture du fichier JSON")
                        # Afficher les 100 premiers caractères pour le débogage
                        content_preview = json_content[:100].decode('utf-8', errors='ignore') if json_content else "EMPTY"
                        return jsonify({
                            "error": "JSON read error", 
                            "details": f"Erreur lors de la lecture du fichier JSON: {str(e)}",
                            "content_preview": content_preview
                        }), 400
                    
                    # Vérification de la structure du JSON
                    if not isinstance(proj_data, dict):
                        logging.error(f"JSON invalide: doit être un objet, reçu: {type(proj_data)}")
                        return jsonify({
                            "error": "Invalid JSON structure", 
                            "details": f"Les données JSON doivent être un objet, type reçu: {type(proj_data)}"
                        }), 400
                    
                    if 'blocks' not in proj_data:
                        logging.warning("Clé 'blocks' manquante dans les données du projet")
                        proj_data['blocks'] = []
                    
                    logging.info(f"Données du projet: {len(proj_data.get('blocks', []))} blocs, types: {proj_data.get('blockTypes', [])}")
                    
                    # Sauvegarder les données
                    save_data(proj_data)
                    logging.info("Données du projet sauvegardées")

                    # ----- extraire les images --------------------------------
                    img_members = [n for n in file_list
                                   if n.startswith('images/') and not n.endswith('/')]
                    logging.info(f"Nombre d'images trouvées: {len(img_members)}")
                    
                    images_folder = None
                    images_info = []
                    if img_members:
                        folder_id     = f"webtoon_{uuid.uuid4().hex[:8]}"
                        images_folder = os.path.join(app.config['UPLOAD_FOLDER'], folder_id)
                        os.makedirs(images_folder, exist_ok=True)
                        logging.info(f"Dossier d'images créé: {images_folder}")
                        
                        for i, member in enumerate(img_members):
                            dest = os.path.join(images_folder,
                                                os.path.relpath(member, 'images'))
                            os.makedirs(os.path.dirname(dest), exist_ok=True)
                            logging.info(f"Extraction de {member} vers {dest}")
                            try:
                                with z.open(member) as src, open(dest, 'wb') as dst:
                                    shutil.copyfileobj(src, dst)
                                
                                # Ajouter les informations d'image pour le frontend
                                filename = os.path.basename(dest)
                                folder = os.path.basename(images_folder)
                                url_path = url_for('serve_upload', fname=f'{folder}/{filename}')
                                
                                images_info.append({
                                    'filename': filename,
                                    'url': url_path,
                                    'path': dest
                                })
                            except Exception as e:
                                logging.error(f"Erreur lors de l'extraction de {member}: {str(e)}")
                                # Continue avec les autres images

                    logging.info("=== Import ZIP réussi ===")
                    logging.info(f"Images extraites: {len(images_info)}")
                    return jsonify({
                        "status": "success",
                        "imagesFolder": os.path.basename(images_folder) if images_folder else None,
                        "images": images_info,  # Ajouter la liste des images
                        "project": proj_data
                    })

            except zipfile.BadZipFile:
                logging.error("Format de fichier ZIP invalide")
                return jsonify({
                    "error": "Bad ZIP format", 
                    "details": "Le fichier n'est pas un ZIP valide. Vérifiez que le fichier n'est pas corrompu."
                }), 400

    except Exception as e:
        logging.exception("Import ZIP failed")
        return jsonify({
            "error": "Server error", 
            "details": str(e),
            "type": type(e).__name__
        }), 500

@app.route('/api/ocr-translate', methods=['POST'])
def api_ocr_translate():
    """Route pour OCR + traduction d'images"""
    if 'image' not in request.files:
        return jsonify({"error": "No image field"}), 400
    
    img_file = request.files['image']
    if img_file.filename == '':
        return jsonify({"error": "Empty filename"}), 400

    try:
        # Ouvrir et redimensionner l'image si nécessaire
        img = Image.open(img_file.stream).convert("RGB")
        
        # Redimensionner si l'image est trop grande (optimisation)
        max_size = (2000, 2000)
        if img.size[0] > max_size[0] or img.size[1] > max_size[1]:
            img.thumbnail(max_size, Image.Resampling.LANCZOS)
            logging.info(f"Image redimensionnée à {img.size}")
        
        # 1) OCR (auto-langue)
        try:
            # Spécifier le chemin de Tesseract pour Windows
            tesseract_path = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
            if os.path.exists(tesseract_path):
                pytesseract.pytesseract.tesseract_cmd = tesseract_path
            
            raw_text = pytesseract.image_to_string(img, lang="eng+fra+jpn+kor")
            raw_text = raw_text.strip()
            if not raw_text:
                return jsonify({"error": "OCR found no text"})
        except Exception as ocr_error:
            logging.error(f"OCR error: {str(ocr_error)}")
            return jsonify({"error": "OCR failed", "details": str(ocr_error)}), 500
        
        # 2) Détection langue pour info
        try:
            detected = langdetect.detect(raw_text)
        except Exception:
            detected = "unknown"

        # 3) Traduction via Gemini
        api_key = load_data().get('api_keys', {}).get('gemini', '')
        if not api_key:
            return jsonify({"error": "Missing Gemini key"})
        
        prompt = f"""Traduits en français.
Donne 3 variantes courtes, sans parenthèses ni explication :
« {raw_text} »"""
        
        try:
            r = requests.post(
                f'https://generativelanguage.googleapis.com/v1beta/models/gemini-1.5-flash:generateContent?key={api_key}',
                json={
                    "contents": [{"parts": [{"text": prompt}]}],
                    "generationConfig": {"temperature": 0.8, "candidateCount": 1}
                },
                timeout=30
            )
            
            if r.status_code != 200:
                return jsonify({"error": "Gemini error", "details": r.text}), 502

            txt = r.json()['candidates'][0]['content']['parts'][0]['text']
            sug = [l.strip().lstrip("•-0123456789. ").split('(', 1)[0].strip()
                   for l in txt.split('\n') if l.strip()]
            
            return jsonify({
                "status": "success",
                "sourceLang": detected,
                "ocr": raw_text,
                "suggestions": sug[:3]
            })
        except Exception as gemini_error:
            logging.error(f"Gemini API error: {str(gemini_error)}")
            return jsonify({"error": "Translation failed", "details": str(gemini_error)}), 500
            
    except Exception as e:
        logging.exception("OCR-Translate failed")
        return jsonify({"error": str(e)}), 500

@app.route('/api/reformulate', methods=['POST'])
def api_reformulate():
    """Route pour reformulation intelligente avec intention spécifique"""
    data = request.json
    if data is None:
        return jsonify({"error": "Invalid JSON"}), 400
    
    api_key = data.get('api_key', '')
    content = data.get('content', '')
    intention = data.get('intention', '')
    custom_intention = data.get('custom_intention', '')
    
    if not api_key or not content:
        return jsonify({"error": "Missing API key or content"}), 400
    
    # ─── PROMPTS SANS INVENTION ──────────────────────────────────────────────
    if intention == "expand":
        prompt = f"""
Tu es un assistant de réécriture.

◆ Objectif : étoffer la phrase ci-dessous **sans** ajouter d'information nouvelle.
◆ Conserve temps, point de vue et registre d'origine.

Règles :
1. N'ajoute aucun nom propre, date, lieu ou évènement absent du texte.
2. N'invente pas de contexte ni de justification implicite.
3. Allonge la phrase d'environ +50 % (max +100 %).
4. Retourne une seule phrase (pas de puces, pas de guillemets).
5. Français soigné, même ton.

Phrase : {content}

Donne **3 variantes** développées, **une par ligne** :
"""

    elif intention == "summarize":
        prompt = f"""
Tu es un assistant de réécriture.

◆ Objectif : condenser la phrase ci-dessous **sans rien omettre d'essentiel**.

Règles :
1. Ne supprime aucune information factuelle.
2. N'ajoute aucun détail nouveau.
3. Réduis la longueur d'environ 40 % à 60 %.
4. Une seule phrase, même temps/verbe et même ton.
5. Pas de puces, pas de guillemets.

Phrase : {content}

Donne **3 variantes** condensées, **une par ligne** :
"""

    elif intention == "rephrase":
        prompt = f"""
Tu es un assistant de réécriture.

◆ Objectif : reformuler la phrase ci-dessous en gardant **exactement le même sens** et
une longueur équivalente (±10 %).

Règles :
1. N'ajoute ni ne retire aucune information.
2. Changements permis : vocabulaire, rythme, structure syntaxique.
3. Longueur finale entre 90 % et 110 % de l'original.
4. Une seule phrase, même registre.
5. Pas de puces, pas de guillemets.

Phrase : {content}

Donne **3 variantes** reformulées, **une par ligne** :
"""

    elif intention == "custom" and custom_intention.strip():
        prompt = f"""
Tu es un assistant de réécriture.

◆ La demande utilisateur est : « {custom_intention} »

Avant d'écrire, **réinterprète mentalement** la demande pour qu'elle respecte :
• aucune information ajoutée ou supprimée  
• même temps/verbe et même point de vue que l'original

Ensuite, applique-la à la phrase ci-dessous.

Phrase : {content}

Donne **3 variantes** conformes à la demande, **une par ligne** :
"""

    else:
        # Fallback très strict
        prompt = f"""
Reformule la phrase ci-dessous sans changer son sens ni sa longueur.
Ne pas ajouter d'élément nouveau.

Phrase : {content}

3 variantes, une par ligne :
"""
    # ─────────────────────────────────────────────────────────────────────────
    
    try:
        response = requests.post(
            f'https://generativelanguage.googleapis.com/v1beta/models/gemini-1.5-flash:generateContent?key={api_key}',
            json={
                'contents': [{'parts': [{'text': prompt}]}],
                'generationConfig': {
                    'temperature': 0.2,          # moins créatif → moins d'inventions
                    'candidateCount': 1,         # une seule sortie, plus stable
                    'top_p': 0.9                 # limite le vocabulaire exotique
                }
            },
            timeout=30
        )
        
        if response.status_code != 200:
            return jsonify({"error": response.json().get('error', {}).get('message', response.text)}), 502
        
        data = response.json()
        raw_text = data.get('candidates', [{}])[0].get('content', {}).get('parts', [{}])[0].get('text', '')
        
        suggestions = []
        for line in raw_text.split('\n'):
            line = line.strip()
            if line:
                # Remove bullets, numbers, and explanations in parentheses
                line = line.lstrip('•-0123456789. \t')
                line = line.split('(', 1)[0].strip()
                if line:
                    suggestions.append(line)
        
        return jsonify({
            "status": "success",
            "suggestions": suggestions[:3],
            "intention": intention,
            "custom_intention": custom_intention
        })
        
    except Exception as e:
        logging.error(f"Reformulation API error: {str(e)}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/upload-reader', methods=['POST'])
def api_upload_reader():
    """Route pour uploader des fichiers dans le lecteur (PDF et images)"""
    if 'file' not in request.files:
        return jsonify({"error": "No file part"}), 400
        
    file = request.files['file']
    
    if file.filename == '':
        return jsonify({"error": "No selected file"}), 400
    
    # Vérifier le type de fichier
    allowed_extensions = {'.pdf', '.png', '.jpg', '.jpeg', '.gif', '.webp', '.bmp'}
    if file.filename:
        file_ext = os.path.splitext(file.filename)[1].lower()
    else:
        return jsonify({"error": "Nom de fichier invalide"}), 400
    
    if file_ext not in allowed_extensions:
        return jsonify({"error": f"Type de fichier non supporté. Extensions autorisées: {', '.join(allowed_extensions)}"}), 400
        
    if file and file.filename:
        # Create a unique filename
        filename = secure_filename(file.filename)
        unique_filename = f"reader_{uuid.uuid4()}{file_ext}"
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], unique_filename)
        
        # Save the file
        file.save(file_path)
        
        # Return the path to access the file
        file_url = url_for('serve_upload', fname=unique_filename)
        
        return jsonify({
            "status": "success",
            "filename": filename,
            "path": file_path,
            "url": file_url,
            "type": "pdf" if file_ext == '.pdf' else "image"
        })
    
    return jsonify({"error": "Failed to upload file"}), 500

@app.route('/api/apply-correction', methods=['POST'])
def api_apply_correction():
    data = request.json or {}
    idx   = data.get('block_index')
    start = data.get('error_start')
    end   = data.get('error_end')
    repl  = data.get('replacement')

    if None in (idx, start, end, repl):
        return jsonify({"error": "Missing required parameters"}), 400

    # Validation des types après vérification de None
    if not isinstance(idx, int) or not isinstance(start, int) or not isinstance(end, int) or not isinstance(repl, str):
        return jsonify({"error": "Invalid parameter types"}), 400

    project = load_data()
    blocks  = project.get('blocks', [])
    if idx >= len(blocks):
        return jsonify({"error": "Invalid block index"}), 400

    orig = blocks[idx]['content']
    if not (0 <= start < end <= len(orig)):
        return jsonify({"error": "Invalid error positions"}), 400

    # On découpe en trois
    before = orig[:start]
    after  = orig[end:]
    repl   = repl.strip()

    # On ne met un espace que s'il manque vraiment :
    sep1 = '' if before.endswith((' ', '\n', '')) or repl.startswith(' ') else ' '
    sep2 = '' if after.startswith((' ', '\n', '')) or repl.endswith(' ') else ' '

    corrected = f"{before}{sep1}{repl}{sep2}{after}"

    # Mise à jour & sauvegarde
    blocks[idx]['content'] = corrected
    save_data(project)

    return jsonify({
        "status":    "success",
        "original":  orig,
        "corrected": corrected,
        "block_index": idx
    })

if __name__ == '__main__':
    app.run(debug=True) 